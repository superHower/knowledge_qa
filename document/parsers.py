"""
文档解析器
"""

import re
from pathlib import Path

from .base import BaseDocumentParser, DocumentContent


class TextParser(BaseDocumentParser):
    """TXT/Markdown 文本解析器"""
    
    async def parse(self, file_path: str) -> DocumentContent:
        """解析文本文件"""
        content = await self.extract_text(file_path)
        
        # 提取标题（第一行或 # 开头的行）
        title = None
        lines = content.split("\n")
        for line in lines[:10]:  # 只看前10行
            line = line.strip()
            if line.startswith("# "):
                title = line[2:].strip()
                break
            elif line and not title:
                title = line[:100]  # 取第一行作为标题
        
        return DocumentContent(
            title=title,
            content=content,
            metadata={"source": file_path},
        )
    
    async def extract_text(self, file_path: str) -> str:
        """提取纯文本，支持多种编码"""
        path = Path(file_path)
        
        # 尝试多种编码
        encodings = ["utf-8", "utf-16", "utf-16-le", "utf-16-be", "gbk", "gb2312", "gb18030", "latin-1"]
        
        for encoding in encodings:
            try:
                content = path.read_text(encoding=encoding)
                return content
            except (UnicodeDecodeError, UnicodeError):
                continue
        
        # 最后尝试二进制读取并解码
        raw_bytes = path.read_bytes()
        # 去掉 BOM
        if raw_bytes.startswith(b'\xff\xfe') or raw_bytes.startswith(b'\xfe\xff'):
            raw_bytes = raw_bytes[2:]
        return raw_bytes.decode("utf-8", errors="replace")


class PDFParser(BaseDocumentParser):
    """PDF 解析器"""
    
    def __init__(self):
        self._pypdf = None
    
    async def parse(self, file_path: str) -> DocumentContent:
        """解析 PDF 文件"""
        from pypdf import PdfReader
        
        content = await self.extract_text(file_path)
        
        # 尝试读取 PDF 元信息
        reader = PdfReader(file_path)
        metadata = {}
        
        if reader.metadata:
            metadata = {
                "title": reader.metadata.get("/Title", ""),
                "author": reader.metadata.get("/Author", ""),
                "subject": reader.metadata.get("/Subject", ""),
                "page_count": len(reader.pages),
            }
        
        return DocumentContent(
            title=metadata.get("title") or Path(file_path).stem,
            content=content,
            metadata=metadata,
        )
    
    async def extract_text(self, file_path: str) -> str:
        """提取 PDF 文本"""
        from pypdf import PdfReader
        
        reader = PdfReader(file_path)
        text_parts = []
        
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        
        return "\n\n".join(text_parts)


class DocxParser(BaseDocumentParser):
    """Word 文档解析器"""
    
    async def parse(self, file_path: str) -> DocumentContent:
        """解析 Word 文档"""
        from docx import Document
        
        doc = Document(file_path)
        
        # 提取标题和内容
        title = None
        content_parts = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # 标题样式判断
            if para.style.name.startswith("Heading"):
                if not title:
                    title = text
                content_parts.append(f"## {text}")
            else:
                content_parts.append(text)
        
        return DocumentContent(
            title=title or Path(file_path).stem,
            content="\n\n".join(content_parts),
            metadata={
                "author": doc.core_properties.author,
                "page_count": len(doc.paragraphs),
            },
        )
    
    async def extract_text(self, file_path: str) -> str:
        """提取 Word 文档文本"""
        from docx import Document
        
        doc = Document(file_path)
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        return "\n\n".join(text_parts)


class HTMLParser(BaseDocumentParser):
    """HTML 解析器"""
    
    async def parse(self, file_path: str) -> DocumentContent:
        """解析 HTML 文件"""
        content = await self.extract_text(file_path)
        
        return DocumentContent(
            title=Path(file_path).stem,
            content=content,
            metadata={"source": file_path},
        )
    
    async def extract_text(self, file_path: str) -> str:
        """提取 HTML 纯文本"""
        import html2text
        
        html_content = Path(file_path).read_text(encoding="utf-8")
        h = html2text.HTML2Text()
        h.ignore_links = False
        h.ignore_images = True
        
        return h.handle(html_content)


class CSVParser(BaseDocumentParser):
    """CSV 解析器"""
    
    async def parse(self, file_path: str) -> DocumentContent:
        """解析 CSV 文件"""
        import pandas as pd
        
        df = pd.read_csv(file_path)
        content = await self.extract_text(file_path)
        
        return DocumentContent(
            title=Path(file_path).stem,
            content=content,
            metadata={
                "columns": list(df.columns),
                "row_count": len(df),
            },
        )
    
    async def extract_text(self, file_path: str) -> str:
        """提取 CSV 文本"""
        import pandas as pd
        
        df = pd.read_csv(file_path)
        text_parts = []
        
        # 添加列名
        text_parts.append(" | ".join(df.columns))
        text_parts.append(" | ".join(["---"] * len(df.columns)))
        
        # 添加数据行
        for _, row in df.head(100).iterrows():  # 限制前100行
            text_parts.append(" | ".join(str(v) for v in row.values))
        
        if len(df) > 100:
            text_parts.append(f"... (共 {len(df)} 行)")
        
        return "\n".join(text_parts)


class ParserFactory:
    """解析器工厂"""
    
    _parsers = {
        ".txt": TextParser,
        ".md": TextParser,
        ".pdf": PDFParser,
        ".docx": DocxParser,
        ".doc": DocxParser,
        ".html": HTMLParser,
        ".htm": HTMLParser,
        ".csv": CSVParser,
    }
    
    @classmethod
    def get_parser(cls, file_path: str) -> BaseDocumentParser | None:
        """获取对应的解析器"""
        ext = Path(file_path).suffix.lower()
        parser_class = cls._parsers.get(ext)
        
        if parser_class:
            return parser_class()
        
        return None
    
    @classmethod
    def register_parser(cls, extension: str, parser_class: type[BaseDocumentParser]):
        """注册新的解析器"""
        cls._parsers[extension.lower()] = parser_class
    
    @classmethod
    def supported_extensions(cls) -> list[str]:
        """获取支持的扩展名"""
        return list(cls._parsers.keys())
